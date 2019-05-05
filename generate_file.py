# -*- coding: utf-8 -*-
import os
if not os.environ.get("DJANGO_SETTINGS_MODULE"):
    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "ResultManage.settings")
import django
django.setup()
import time
import random
import subprocess
import docx
from datetime import datetime
import re
# import psycopg2
import sqlite3
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION


def generate_docx(handoutlist, fileinfo_list, templates_dir, handoutlist_docxs, handoutlist_docs):
    result_count = len(fileinfo_list)
    template_filepath = os.path.join(templates_dir, "template" + str(result_count) + ".docx")
    print(template_filepath)
    doc = docx.Document(template_filepath)
    # 如果只为null,替换成空字符串""
    handoutlist_propertys = [
        handoutlist.title,
        handoutlist.listnum,
        handoutlist.signer,
        handoutlist.name,
        handoutlist.purpose,
        handoutlist.auditnum,
        handoutlist.secrecyagreementnum,
        handoutlist.signature,
        handoutlist.medianums,
        handoutlist.sendunit,
        handoutlist.receiveunit,
        handoutlist.sendunitaddr,
        handoutlist.receiveunitaddr,
        handoutlist.sendunitpostcode,
        handoutlist.receiveunitpostcode,
        handoutlist.handler,
        handoutlist.receiver,
        handoutlist.handlerphonenum,
        handoutlist.receiverphonenum,
        handoutlist.handlermobilephonenum,
        handoutlist.receivermobilephonenum,
        handoutlist.sendouttime,
        handoutlist.recievetime,
        handoutlist.mapnums,
        handoutlist.filename,
        handoutlist.file,
    ]
    # 防止有的值为null出错
    for handoutlist_property in handoutlist_propertys:
        if handoutlist_property is None or handoutlist_property == "None":
            handoutlist_propertys[handoutlist_propertys.index(handoutlist_property)] = ""

    # print(len(doc.paragraphs))
    # 清单名称
    title = handoutlist_propertys[0]
    # doc.paragraphs[0].runs[0].text = title
    title_runs = doc.paragraphs[0].runs
    for run in title_runs:
        if title_runs.index(run) == 0:
            run.text = title
        else:
            run.text = ""
    # 编号
    listnum = handoutlist_propertys[1]
    signer = handoutlist_propertys[2]
    # 签发人
    for run in doc.paragraphs[1].runs:
        if run.text.strip() == "listnum":
            run.text = listnum
            run.underline = True
        if run.text.strip() == "signer":
            run.text = signer
            run.underline = True

    # 名称
    name = handoutlist_propertys[3]
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace("name", name)

    # 用途
    perpose = handoutlist_propertys[4]
    doc.paragraphs[4].text = doc.paragraphs[4].text.replace("支援保障", perpose)
    # 审单
    auditnum = handoutlist_propertys[5]
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("auditnum", auditnum)
    # 协议编号
    secrecyagreementnum = handoutlist_propertys[6]
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("secrecyagreementnum", secrecyagreementnum)

    # 递送方式, 递送签名
    # selfgetway  postway  networkway  sendtoway
    selfgetway = handoutlist.selfgetway
    postway = handoutlist.postway
    networkway = handoutlist.networkway
    sendtoway = handoutlist.sendtoway
    signature = handoutlist_propertys[7]
    way1 = "☑自取" if selfgetway else "□自取"
    way2 = "☑邮寄" if postway else "□邮寄"
    way3 = "☑网络" if networkway else "□网络"
    way4 = "☑送往" if sendtoway else "□送往"
    ways = way1 + "  " + way2 + "  " + way3 + "  " + way4
    # doc.paragraphs[6].text = doc.paragraphs[6].text.replace("□自取  □邮寄  □网络  □送往 （签名）signature",ways)
    doc.paragraphs[6].text = "递送方式： " + ways + " （签名）"
    run6 = doc.paragraphs[6].add_run(signature)
    run6.underline = True

    # 介质类型
    papermedia = handoutlist.papermedia
    cdmedia = handoutlist.cdmedia
    diskmedia = handoutlist.diskmedia
    networkmedia = handoutlist.networkmedia
    othermedia = handoutlist.othermedia
    media1 = "☑纸质" if papermedia else "□纸质"
    media2 = "☑光盘" if cdmedia else "□光盘"
    media3 = "☑硬盘" if diskmedia else "□硬盘"
    media4 = "☑网络" if networkmedia else "□网络"
    media5 = "☑其他" if othermedia else "□其他"
    # 介质类型
    medias = media1 + "  " + media2 + "  " + media3 + "  " + media4 + "  " + media5
    # 介质编号
    medianums = handoutlist_propertys[8]
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("□纸质  □光盘  □硬盘  □网络  □其他", medias)
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("medianums", medianums)
    doc.paragraphs[7].text ="介质类型： " + medias+"      介质编号："+medianums

    # 发出单位
    sendunit = handoutlist_propertys[9]
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("sendunit", sendunit)
    # 接收单位
    receiveunit = handoutlist_propertys[10]
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("receiveunit", receiveunit)

    # 发出单位通讯地址
    sendunitaddr = handoutlist_propertys[11]
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("sendunitaddr", sendunitaddr)
    # 接收单位通讯地址
    receiveunitaddr = handoutlist_propertys[12]
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("receiveunitaddr", receiveunitaddr)
    # 发出单位邮政编码
    sendunitpostcode = handoutlist_propertys[13]
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("sendunitpostcode", sendunitpostcode)
    # 接收单位邮政编码
    receiveunitpostcode = handoutlist_propertys[14]
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("receiveunitpostcode", receiveunitpostcode)

    # 经办人
    handler = handoutlist_propertys[15]
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("handler", handler)
    # 接收人
    receiver = handoutlist_propertys[16]
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("receiver", receiver)
    # 经办人联系电话(座机)
    handlerphonenum = handoutlist_propertys[17]
    # if handlerphonenum is None:
    #     handlerphonenum =""
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("handlerphonenum", handlerphonenum)
    # 接收人联系电话(座机)
    receiverphonenum = handoutlist_propertys[18]
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("receiverphonenum", receiverphonenum)
    # 经办人联系电话(手机)
    handlermobilephonenum = handoutlist_propertys[19]
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("handlermobilephonenum", handlermobilephonenum)
    # 接收人联系电话(手机)
    receivermobilephonenum = handoutlist_propertys[20]
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("receivermobilephonenum", receivermobilephonenum)

    # 发出日期
    sendouttime = handoutlist_propertys[21]
    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("sendouttime", sendouttime)
    # 接收日期
    recievetime = handoutlist_propertys[22]
    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("recievetime", recievetime)

    table = doc.tables[0]
    count = 1
    for fileinfo in fileinfo_list:
        row = table.rows[count]
        # row_cells = row.cells
        # for cell in row_cells:
        #     # 表格内容居中
        #     # cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
        #     if row_cells.index(cell) == 1:
        #         cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平左对齐
        #     # else:
        #     #     cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
        #     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直对齐
        # print(len(row.cells[1].paragraphs[0].runs))

        row.cells[0].paragraphs[0].runs[0].text = str(count)
        # name
        row.cells[1].paragraphs[0].runs[0].text = fileinfo.name if fileinfo.name is not None and fileinfo.name != "None" else ""
        # secretlevel
        row.cells[2].paragraphs[0].runs[0].text = fileinfo.secretlevel if fileinfo.secretlevel is not None and fileinfo.secretlevel != "None" else ""
        # resultnum
        row.cells[3].paragraphs[0].runs[0].text = fileinfo.resultnum if fileinfo.resultnum is not None and fileinfo.resultnum != "None" else ""
        # datasize
        row.cells[4].paragraphs[0].runs[0].text = fileinfo.datasize if fileinfo.datasize is not None and fileinfo.datasize != "None" else ""
        # formatormedium
        row.cells[5].paragraphs[0].runs[0].text = fileinfo.formatormedium if fileinfo.formatormedium is not None and fileinfo.formatormedium != "None" else ""
        # remarks
        row.cells[6].paragraphs[0].runs[0].text = fileinfo.remarks if fileinfo.remarks is not None and fileinfo.remarks != "None" else ""
        count += 1

    mapnums = handoutlist_propertys[23]
    # doc.paragraphs[18].text = doc.paragraphs[18].text.replace("mapnums", mapnums)
    doc.paragraphs[18].text = doc.paragraphs[18].text + mapnums
    # 添加段落
    # doc.add_paragraph(mapnums)
    # filename = handoutlist[34]
    filename = "测绘" + datetime.now().strftime("%Y") + "-" + "%04d" % handoutlist.id + ".docx"
    docx_filepath = os.path.join(handoutlist_docxs, filename)
    doc.save(docx_filepath)

    handoutlist.file = os.path.join("handoutlist_docxs",filename)
    print("调用了generate_docx函数")
    return handoutlist

    # 转换为doc格式
    # subprocess.check_output(
    #     ["soffice", "--headless", "--invisible", "--convert-to", "doc", docx_filepath, "--outdir", handoutlist_docs])


if __name__ == '__main__':
    # generate_docx(handoutlist, fileinfo_list, templates_dir, handoutlist_docxs, handoutlist_docs)
    pass