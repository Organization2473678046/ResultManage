#!/opt/rh/rh-python36/root/usr/bin/python3
# -*- coding: utf-8 -*-
import os
import sys
import psycopg2
from celery_app import app
from datetime import datetime
import sqlite3
from docx import Document
import re
import subprocess
import time
from django.db import transaction


@app.task
def doc_data_updata(dbname,filepath, handoutlist_uniquenum, filename):

    print(filepath)
    print(handoutlist_uniquenum)
    print(filename)
    if filename.split(".").pop() == "doc":
        path_list = filepath.split("/")
        path_list.pop()
        out_dir = "/".join(path_list)
        out_format = "docx"
        out_put = subprocess.check_output(
            ["soffice", "--headless", "--invisible", "--convert-to", out_format, filepath, "--outdir", out_dir])
        # get_doc_data
        # print(out_put,"**********************")
        document = Document(filepath + "x")
    else:
        document = Document(filepath)

    title = ""      # 分发单标题
    signer = ""     # 签发
    name = ""       # 分发单名称
    uniquenum = handoutlist_uniquenum        # 分发单唯一编号
    listnum = ""    # 编号
    auditnum = ""   # 审核编号
    secrecyagreementnum = ""    # 保密责任书编号
    purpose = ""    # 用途
    sendunit = ""   # 发出单位
    sendunitaddr = ""       # 发出单位通讯地址
    sendunitpostcode = ""   # 发出单位邮政编码
    receiveunit = ""        # 接收单位
    receiveunitaddr = ""    # 接收单位通讯地址
    receiveunitpostcode = ""    # 接收单位邮政编码
    handler = ""            # 经办人
    handlerphonenum = ""    # 经办人联系电话(座机)
    handlermobilephonenum = ""  # 经办人联系电话(手机)
    receiver = ""           # 接收人
    receiverphonenum = ""   # 接收人联系电话(座机)
    receivermobilephonenum = ""     # 接收人联系电话(手机)
    sendouttime = ""        # 发出日期
    recievetime = ""        # 接收日期
    selfgetway = False         # 递送方式为自取
    postway = False            # 递送方式为邮寄
    networkway = False         # 递送方式为网络
    sendtoway = False          # 递送方式为送往
    signature = ""          # 递送方式后面的签名
    papermedia = False         # 纸质介质
    cdmedia = False            # 光盘介质
    diskmedia = False          # 硬盘介质
    networkmedia = False       # 网络介质
    othermedia = False         # 其他介质
    medianums = ""          # 介质编号
    mapnums = ""            # 分发单包含的成果的图幅号
    filename = filename           # 分发单文件名字
    file = filepath               # 分发单文件路径

    remake_str = ""
    is_remake = False
    is_textbox = True
    conn = psycopg2.connect(dbname=dbname,
                            user="postgres",
                            password="Lantucx2018",
                            host="localhost",
                            port="5432")
    cursor = conn.cursor()

    x = 0

    # 读取文本框
    xml = document._element.xml
    pattern = "<v:textbox([\s\S]*?)</v:textbox>"
    regex = re.compile(pattern)
    textbox_list = regex.findall(xml)
    if len(textbox_list) <= 2 and len(textbox_list) > 0:
        for result in textbox_list:
            # pattern2 = u"[\u4e00-\u9fa5]+"
            pattern2 = "<w:t>(.*?)</w:t>"
            regex2 = re.compile(pattern2)
            textbox_str = regex2.findall(result)
            textbox_str = "".join(textbox_str)
            if x == 0 and is_textbox == True:
                # print("名称:" + textbox_str)
                name = textbox_str
            if x == 1 and is_textbox == True:
                # print("用途:" + textbox_str)
                purpose = textbox_str
                x = 0
                is_textbox = False
            x += 1
    elif len(textbox_list) == 0:
        for para in document.paragraphs:
            pattern = "[\u4e00-\u9fa5]+"
            regex = re.compile(pattern)
            results = regex.findall(para.text)
            if results != [] and is_remake is False:
                if "称" in results and "名" in results:
                    name = para.text.split("：").pop()
                if "用" in results and "途" in results:
                    purpose = para.text.split("：").pop()
    else:
        print("此文件为多单号文件：" + filename)
        return



    # 读取文本
    for para in document.paragraphs:
        pattern = "[\u4e00-\u9fa5]+"
        regex = re.compile(pattern)
        results = regex.findall(para.text)
        # print(results)
        if results != [] and is_remake is False:
            if "交付清单" in results[0]:
                title = results[0]
                # print("标题：" + results[0])

            if "审批单号" in results:
                my_num01 = para.text.split(" ")[0].split("：").pop()
                auditnum = my_num01
                # print("审批单号:" + my_num01)

            if "签发" in results:
                my_num001 = para.text.split("：").pop().split("签")[0].strip()
                listnum = my_num001
                # print("编号:" + my_num001)

            if "保密责任证书编号" in results:
                my_num02 = para.text.split("：").pop()
                secrecyagreementnum = my_num02
                # print("保密责任证书:" + my_num02)

            if "递送方式" in results:
                my_list01 = para.text.split(" ")
                for my_way in my_list01:
                    if my_way == "自取":
                        selfgetway = True
                    if my_way == "邮寄":
                        postway = True
                    if my_way == "网络":
                        networkway = True
                    if my_way == "送往":
                        sendtoway = True

            if "介质类型" in results:
                my_list02 = para.text.split(" ")
                for my_way in my_list02:
                    if my_way == "纸质":
                        selfgetway = True
                    if my_way == "光盘":
                        postway = True
                    if my_way == "硬盘":
                        networkway = True
                    if my_way == "网络":
                        sendtoway = True
                    if my_way == "其他":
                        sendtoway = True

            if "介质编号" in results:
                my_num03 = para.text.split("：").pop()
                medianums = my_num03
                # print("介质编号:" + my_num03)

            if "发出单位" in results:
                my_str01 = para.text.split(" ")[0].split("：").pop()
                sendunit = my_str01
                # print("发出单位:" + my_str01)

            if "接收单位" in results:
                my_str02 = para.text.split(" ").pop().split("：").pop()
                receiveunit = my_str02
                # print("接收单位：" + my_str02)

            if "通讯地址" in results:
                my_str03 = para.text.split(" ")[0].split("：").pop()
                sendunitaddr = my_str03
                # print("发出方通讯地址：" + my_str03)
                my_str03_ = para.text.split("：").pop()
                receiveunitaddr = my_str03_
                # print("接收方通讯地址：" + receiveunitaddr)


            if "邮政编码" in results:
                my_str04 = para.text.split(" ")[0].split("：").pop()
                sendunitpostcode = my_str04
                # print("发出方邮政编码：" + my_str04)

                my_str04_ = para.text.split("：").pop()
                receiveunitpostcode = my_str04_
                # print("接收方邮政编码：" + receiveunitpostcode)

            if "座机" in results:
                my_str05 = para.text.split(" ")[0].split("）").pop()
                handlerphonenum = my_str05
                # print("发出方联系电话：（座机）：" + my_str05)

                my_str05_ = para.text.split("）").pop()
                receiverphonenum = my_str05_
                # print("接收方联系电话：（座机）：" + receiverphonenum)

            if "手机" in results:
                my_str06 = para.text.split("）")[1].split(" ")[0]
                handlermobilephonenum = my_str06
                # print("发出方联系电话：（手机）：" + my_str06)

                my_str06_ = para.text.split("）").pop()
                receivermobilephonenum = my_str06_
                # print("接收方联系电话：（手机）：" + receivermobilephonenum)

            if "发出日期" in results:
                my_str07 = para.text.split("：")[1].split("接")[0]
                my_str07 = my_str07.strip()
                my_str07_year = my_str07.split("年")[0].strip() if my_str07.split("年")[0].strip() != '' else "  "
                my_str07_mon = my_str07.split("年")[1].split("月")[0].strip() if my_str07.split("年")[1].split("月")[0].strip() != '' else "  "
                my_str07_day = my_str07.split("年")[1].split("月")[1].split("日")[0].strip() if my_str07.split("年")[1].split("月")[1].split("日")[0].strip() != '' else "  "
                sendouttime = my_str07_year + "年" + my_str07_mon + "月" + my_str07_day + "日"
                # print("发出日期：" + sendouttime)

                my_str07_ = para.text.split("接收日期：")[1].strip()
                my_str07_year = my_str07_.split("年")[0].strip() if my_str07_.split("年")[0].strip() != '' else "  "
                my_str07_mon = my_str07_.split("年")[1].split("月")[0].strip() if my_str07_.split("年")[1].split("月")[0].strip() != '' else "  "
                my_str07_day = my_str07_.split("年")[1].split("月")[1].split("日")[0].strip() if my_str07_.split("年")[1].split("月")[1].split("日")[0].strip() != '' else "  "

                recievetime = my_str07_year + "年" + my_str07_mon + "月" + my_str07_day + "日"
                # print("接收日期：" + recievetime)


            if "注" in results:
                is_remake = True
                my_str09 = para.text.split("：").pop()
                remake_str = my_str09
                continue

        if is_remake:
            remake_str += para.text


    mapnums = remake_str
    SQL = "UPDATE results_handoutlist SET title = '{0}', signer='{1}', name='{2}', listnum='{3}', auditnum='{4}', secrecyagreementnum='{5}', purpose='{6}', sendunit='{7}', sendunitaddr='{8}', sendunitpostcode='{9}', receiveunit='{10}', receiveunitaddr='{11}', receiveunitpostcode='{12}', handler='{13}', handlerphonenum='{14}', handlermobilephonenum='{15}', receiver='{16}', receiverphonenum='{17}', receivermobilephonenum='{18}', sendouttimec='{19}', recievetimec='{20}', selfgetway='{21}', postway='{22}', networkway='{23}', sendtoway='{24}', signature='{25}', papermedia='{26}', cdmedia='{27}', diskmedia='{28}', networkmedia='{29}', othermedia='{30}', medianums='{31}', mapnums='{32}' WHERE  uniquenum='{33}'".format(title, signer, name, listnum, auditnum, secrecyagreementnum, purpose, sendunit, sendunitaddr, sendunitpostcode, receiveunit, receiveunitaddr, receiveunitpostcode, handler, handlerphonenum, handlermobilephonenum, receiver, receiverphonenum, receivermobilephonenum, sendouttime, recievetime, selfgetway, postway, networkway, sendtoway, signature, papermedia, cdmedia, diskmedia, networkmedia, othermedia, medianums, mapnums, uniquenum)
    cursor.execute(SQL)



    SQL_UPDATE = "UPDATE results_fileinfo SET isdelete=TRUE WHERE handoutlist_uniquenum='{0}'".format(
        handoutlist_uniquenum)
    print(SQL_UPDATE)
    cursor.execute(SQL_UPDATE)

    # 读取表格
    p = 1
    tables = document.tables
    table = tables[0]
    if listnum[:4] == "2017":
        tables = document.tables
        table = tables[0]
        for i in table.rows:
            table_list = []
            for x in i.cells:
                table_list.append(x.text)
            print(table_list)
            name = table_list[1]
            secretlevel = table_list[2]
            resultnum = table_list[3]
            datasize = ""
            formatormedia = table_list[4]
            remarks = table_list[5]
            handoutlist_uniquenum = handoutlist_uniquenum

            if p == 1:
                p = 2
            else:
                SQL_INSERT = "INSERT INTO results_fileinfo (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum, isdelete) VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}', '{6}', '{7}')".format(name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum, False)
                print(SQL_INSERT)
                cursor.execute(SQL_INSERT)
    else:
        for i in table.rows:
            table_list = []
            for x in i.cells:
                table_list.append(x.text)
            name = table_list[1]
            secretlevel = table_list[2]
            resultnum = table_list[3]
            datasize = table_list[4]
            formatormedia = table_list[5]
            remarks = table_list[6]
            handoutlist_uniquenum = handoutlist_uniquenum

            if p == 1:
                p = 2
            else:
                # SQL_UPDATE = "UPDATE results_fileinfo SET isdelete=TRUE WHERE handoutlist_uniquenum='{0}'".format(handoutlist_uniquenum)
                SQL_INSERT = "INSERT INTO results_fileinfo (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum, isdelete) VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}', '{6}', '{7}')".format(name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum, False)
                print(SQL_INSERT)
                cursor.execute(SQL_INSERT)

    conn.commit()
    conn.close()
    return "successful"




if __name__ == '__main__':
    filepath = '/home/ltcx/Desktop/ResultManage/media/data/2019/05/17/2019-05-17-15-18-41-006062/20190517.doc'
    handoutlist_uniquenum = '20190517143616593026_1'
    filename = '20190517.doc'
    doc_data_updata(filepath, handoutlist_uniquenum, filename)

