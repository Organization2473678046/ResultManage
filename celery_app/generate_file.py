# -*- coding: utf-8 -*-
# coding=utf8

import os
import sys

# import io
# sys.stdout = io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
import psycopg2
import time
import random
import subprocess
import docx
from datetime import datetime
import re
import chardet
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
from docx.shared import Pt, Cm
from celery_app import app


@app.task
def generate_docx(dbname, handoutlist_id, handoutlist_uniquenum, templates_dir, handoutlist_docxs):
    handoutlist, fileinfo_list = get_handoutlist_data(dbname, handoutlist_id, handoutlist_uniquenum)
    print(handoutlist)
    print(fileinfo_list)
    template_filepath = os.path.join(templates_dir, "template" + ".docx")
    # TODO
    if not os.path.exists(template_filepath):
        return
    doc = docx.Document(template_filepath)
    # 如果只为null,替换成空字符串""
    handoutlist_propertys = list(handoutlist)
    # 防止有的值为null出错
    for handoutlist_property in handoutlist_propertys:
        if handoutlist_property is None or handoutlist_property == "None":
            handoutlist_propertys[handoutlist_propertys.index(handoutlist_property)] = ""

    # print(len(doc.paragraphs))
    # 清单名称
    title = handoutlist_propertys[0]
    # doc.paragraphs[0].runs[0].text = title
    title_runs = doc.paragraphs[0].runs
    for run in title_runs:
        if title_runs.index(run) == 0:
            run.text = title
        else:
            run.text = ""
    # 编号
    listnum = handoutlist_propertys[1]
    signer = handoutlist_propertys[2]
    # 签发人
    for run in doc.paragraphs[1].runs:
        if run.text.strip() == "listnum":
            run.text = " " + listnum + " "
            run.underline = True
        if run.text.strip() == "signer":
            run.text = "  " + signer + "  "
            run.underline = True

    # 名称
    name = handoutlist_propertys[3]
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace("name", name)

    # 用途
    perpose = handoutlist_propertys[4]
    doc.paragraphs[4].text = doc.paragraphs[4].text.replace("支援保障", perpose)
    # 审单
    auditnum = handoutlist_propertys[5]
    # if re.match(r"测资审〔\d+〕\d+号",auditnum):
    #     auditnum = process_str(auditnum, 32)
    # else:
    auditnum = process_str(auditnum, 35)
    # auditnum = process_str(auditnum, 35-len(re.findall("\d+[\u4E00-\u9FA5]|[a-zA-Z][\u4E00-\u9FA5]",auditnum)))

    # if re.match(r"〔\d*〕",auditnum):
    #     auditnum = process_str(auditnum, 33)
    # elif re.match(r"〔\d*〕", auditnum):
    #     auditnum = process_str(auditnum, 33)
    # else:
    #     auditnum = process_str(auditnum, 35)
    # if len(auditnum) ==11:
    #     auditnum = auditnum+" "*23
    # elif len(auditnum) ==12:
    #     auditnum = auditnum+" "*22
    # elif len(auditnum) == 13:
    #     auditnum = auditnum + " " * 21
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("auditnum", auditnum)
    # 协议编号
    secrecyagreementnum = handoutlist_propertys[6]
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("secrecyagreementnum", secrecyagreementnum)

    # 递送方式, 递送签名
    # selfgetway  postway  networkway  sendtoway
    selfgetway = handoutlist_propertys[7]
    postway = handoutlist_propertys[8]
    networkway = handoutlist_propertys[9]
    sendtoway = handoutlist_propertys[10]
    signature = handoutlist_propertys[11]
    way1 = "☑自取" if selfgetway else "□自取"
    way2 = "☑邮寄" if postway else "□邮寄"
    way3 = "☑网络" if networkway else "□网络"
    way4 = "☑送往" if sendtoway else "□送往"
    ways = way1 + "  " + way2 + "  " + way3 + "  " + way4
    # doc.paragraphs[6].text = doc.paragraphs[6].text.replace("□自取  □邮寄  □网络  □送往 （签名）signature",ways)
    # doc.paragraphs[6].text = "递送方式： " + ways + " （签名）"
    # run6 = doc.paragraphs[6].add_run(signature)
    # run6.underline = True
    print(doc.paragraphs[6].text)

    # 介质类型
    papermedia = handoutlist_propertys[12]
    cdmedia = handoutlist_propertys[13]
    diskmedia = handoutlist_propertys[14]
    networkmedia = handoutlist_propertys[15]
    othermedia = handoutlist_propertys[16]
    media1 = "☑纸质" if papermedia else "□纸质"
    media2 = "☑光盘" if cdmedia else "□光盘"
    media3 = "☑硬盘" if diskmedia else "□硬盘"
    media4 = "☑网络" if networkmedia else "□网络"
    media5 = "☑其他" if othermedia else "□其他"
    # 介质类型
    medias = media1 + "  " + media2 + "  " + media3 + "  " + media4 + "  " + media5
    # 介质编号
    medianums = handoutlist_propertys[17]
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("□纸质  □光盘  □硬盘  □网络  □其他", medias)
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("medianums", medianums)
    # doc.paragraphs[7].text = "介质类型： " + medias + "      介质编号：" + medianums
    doc.paragraphs[7].text = doc.paragraphs[7].text.replace("medianums", medianums)
    for run in doc.paragraphs[7].runs:
        print(run.text)
        if "□" in run.text:
            run.font.size = Pt(10.5)
            run.font.name = "宋体"
    # 发出单位
    sendunit = handoutlist_propertys[18]
    sendunit = process_str(sendunit, 35)
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("sendunit", sendunit)
    # 接收单位
    receiveunit = handoutlist_propertys[19]
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("receiveunit", receiveunit)

    # 发出单位通讯地址
    sendunitaddr = handoutlist_propertys[20]
    sendunitaddr = process_str(sendunitaddr, 35)
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("sendunitaddr", sendunitaddr)
    # 接收单位通讯地址
    receiveunitaddr = handoutlist_propertys[21]
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("receiveunitaddr", receiveunitaddr)
    # 发出单位邮政编码
    sendunitpostcode = handoutlist_propertys[22]
    sendunitpostcode = process_str(sendunitpostcode, 35)
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("sendunitpostcode", sendunitpostcode)
    # 接收单位邮政编码
    receiveunitpostcode = handoutlist_propertys[23]
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("receiveunitpostcode", receiveunitpostcode)

    # 经办人
    handler = handoutlist_propertys[24]
    handler = process_str(handler, 28)
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("handler", handler)
    # 接收人
    receiver = handoutlist_propertys[25]
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("receiver", receiver)
    # 经办人联系电话(座机)
    handlerphonenum = handoutlist_propertys[26]
    handlerphonenum = process_str(handlerphonenum, 28)
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("handlerphonenum", handlerphonenum)
    # 接收人联系电话(座机)
    receiverphonenum = handoutlist_propertys[28]
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("receiverphonenum", receiverphonenum)
    # 经办人联系电话(手机)
    handlermobilephonenum = handoutlist_propertys[28]
    handlermobilephonenum = process_str(handlermobilephonenum, 37)
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("handlermobile", handlermobilephonenum)
    # 接收人联系电话(手机)
    receivermobilephonenum = handoutlist_propertys[29]
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("receivermobile", receivermobilephonenum)

    # 发出日期
    sendouttime = handoutlist_propertys[30]
    if sendouttime != "":
        sendouttime = sendouttime.split("-")[0]+"年"+sendouttime.split("-")[1]+"月"+sendouttime.split("-")[2]+"日"
    else:
        sendouttime = "    年"+"  月"+"  日"
    sendouttime = process_str(sendouttime, 35)

    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("sendouttime", sendouttime)
    # 接收日期
    recievetime = handoutlist_propertys[31]
    if recievetime != "":
        recievetime = recievetime.split("-")[0]+"年"+recievetime.split("-")[1]+"月"+recievetime.split("-")[2]+"日"
    else:
        recievetime = "    年"+"  月"+"  日"
    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("recievetime", recievetime)

    table = doc.tables[0]
    # run_style = table.rows[0].cells[0].paragraphs[0].runs[0].style
    # run_style_font = table.rows[0].cells[0].paragraphs[0].runs[0].style.font
    # table.style="Table Grid"
    # print(table.style,"**********")
    count = 1
    # print(table.cell(0,1).width)
    # print(table.rows[0].cells[0].paragraphs[0].style)
    # print(table.rows[0].cells[0].paragraphs[0].runs[0].style)
    # print(dir(table.rows[0].cells[0].paragraphs[0].runs[0].style))
    # print(table.rows[0].cells[0].paragraphs[0].runs[0].style.font)
    # print(dir(table.rows[0].cells[0].paragraphs[0].runs[0].style.font))

    for fileinfo in fileinfo_list:
        runs = []
        row = table.add_row()
        # row.height = Cm(1)
        row_cells = row.cells
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(8)
            cell.paragraphs[0].paragraph_format.space_before = Pt(8)
            # 表格内容居中
            # cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
            if row_cells.index(cell) == 1:
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平左对齐
            else:
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直对齐

        # TODO
        # row_cells[0].width = Cm(1.3)
        # row_cells[1].width = Cm(7)
        # row_cells[2].width = Cm(1.5)
        # row_cells[3].width = Cm(1.5)
        # row_cells[4].width = Cm(1.81)
        # row_cells[5].width = Cm(2)
        # row_cells[6].width = Cm(1.99)
        run0 = row_cells[0].paragraphs[0].add_run(str(count))
        runs.append(run0)
        run1 = row_cells[1].paragraphs[0].add_run(
            fileinfo[0] if fileinfo[0] is not None and fileinfo[0] != "None" else "")
        runs.append(run1)
        run2 = row_cells[2].paragraphs[0].add_run(
            fileinfo[1] if fileinfo[1] is not None and fileinfo[1] != "None" else "")
        runs.append(run2)
        run3 = row_cells[3].paragraphs[0].add_run(
            fileinfo[2] if fileinfo[2] is not None and fileinfo[1] != "None" else "")
        runs.append(run3)
        run4 = row_cells[4].paragraphs[0].add_run(
            fileinfo[3] if fileinfo[3] is not None and fileinfo[1] != "None" else "")
        runs.append(run4)
        run5 = row_cells[5].paragraphs[0].add_run(
            fileinfo[4] if fileinfo[4] is not None and fileinfo[1] != "None" else "")
        runs.append(run5)
        run6 = row_cells[6].paragraphs[0].add_run(
            fileinfo[5] if fileinfo[5] is not None and fileinfo[1] != "None" else "")
        runs.append(run6)
        for run in runs:
            run.font.size = Pt(9)
            run.font.name = "宋体"
        count += 1

        # row_cells[0].paragraphs[0].text = str(count)
        # name
        # row_cells[1].paragraphs[0].text = fileinfo[0] if fileinfo[0] is not None and fileinfo[0] != "None" else ""
        # row_cells[2].paragraphs[0].text = fileinfo[1] if fileinfo[1] is not None and fileinfo[1] != "None" else ""
        # # resultnum
        # row_cells[3].paragraphs[0].text = fileinfo[2] if fileinfo[2] is not None and fileinfo[2] != "None" else ""
        # # datasize
        # row_cells[4].paragraphs[0].text = fileinfo[3] if fileinfo[3] is not None and fileinfo[3] != "None" else ""
        # # formatormedium
        # row_cells[5].paragraphs[0].text = fileinfo[4] if fileinfo[4] is not None and fileinfo[4] != "None" else ""
        # # remarks
        # row_cells[6].paragraphs[0].text = fileinfo[5] if fileinfo[5] is not None and fileinfo[5] != "None" else ""

    mapnums = handoutlist_propertys[32]
    # doc.paragraphs[18].text = doc.paragraphs[18].text.replace("mapnums", mapnums)
    # doc.paragraphs[18].text = doc.paragraphs[18].text + mapnums
    run_ = doc.paragraphs[19].add_run(mapnums)
    run_.font.size = Pt(9)
    run_.font.name = "宋体"
    # 添加段落
    # doc.add_paragraph(mapnums)
    # filename = handoutlist[35]
    # filename = "测绘"+datetime.now().strftime("%Y-%m-%d-%H-%M-%S") + "-" + "%04d" % handoutlist_propertys[36] + ".docx"
    filename = datetime.now().strftime("%Y-%m-%d-%H-%M-%S") + "-" + "%04d" % handoutlist_propertys[36] + ".docx"

    docx_filepath = os.path.join(handoutlist_docxs, filename)
    doc.save(docx_filepath)

    handoutlist_file = os.path.join("handoutlist_docxs", filename)
    print(handoutlist_file)
    change_postgresql(dbname, handoutlist_id, handoutlist_file)
    return True


def process_str(in_str, length):
    hanzi_regex = re.compile(r'[\u4E00-\u9FA5〔〕]')
    hanzi_list = hanzi_regex.findall(in_str)
    # if length == 35 or length == 28:
    in_str_length = len(hanzi_list) + len(in_str)
    if in_str_length <= length:
        length = length - len(re.findall("\d+[\u4E00-\u9FA5]|[a-zA-Z][\u4E00-\u9FA5]", in_str))
        out_str = in_str + (length - in_str_length) * u" "
        return out_str
    return in_str


def get_handoutlist_data(dbname, handoutlist_id, handoutlist_uniquenum):
    conn = psycopg2.connect(dbname=dbname,
                            user="postgres",
                            password="Lantucx2018",
                            host="localhost",
                            port="5432")
    cursor = conn.cursor()
    SELECT_SQL1 = "select title,listnum,signer,name,purpose,auditnum,secrecyagreementnum,selfgetway, postway,networkway,sendtoway,signature,papermedia,cdmedia, diskmedia, networkmedia, othermedia,medianums,sendunit,receiveunit, sendunitaddr,receiveunitaddr, sendunitpostcode,  receiveunitpostcode, handler,receiver,handlerphonenum,receiverphonenum,handlermobilephonenum, receivermobilephonenum, sendouttime,recievetime,mapnums, filename,file,uniquenum,id from results_handoutlist where uniquenum='%s'" % handoutlist_uniquenum
    cursor.execute(SELECT_SQL1)
    handoutlist = cursor.fetchone()

    SELECT_SQL2 = "select name,secretlevel,resultnum,datasize,formatormedia,remarks from results_fileinfo where handoutlist_uniquenum='%s' order by id" % handoutlist_uniquenum
    cursor.execute(SELECT_SQL2)
    fileinfo_list = cursor.fetchall()
    conn.commit()
    conn.close()
    return handoutlist, fileinfo_list


def change_postgresql(dbname, handoutlist_id, hadoutlist_file):
    conn = psycopg2.connect(dbname=dbname,
                            user="postgres",
                            password="Lantucx2018",
                            host="localhost",
                            port="5432")
    cursor = conn.cursor()
    sql = "update results_handoutlist set file='{0}' where id={1}".format(hadoutlist_file, handoutlist_id)
    cursor.execute(sql)
    conn.commit()
    conn.close()


if __name__ == '__main__':
    generate_docx("resmanagev0.3", 5, "20190506194324094680000001",
                  "/opt/rh/httpd24/root/var/www/html/ResultManage/templates/docx_templates",
                  "/opt/rh/httpd24/root/var/www/html/ResultManage/media/handoutlist_docxs")
    pass
