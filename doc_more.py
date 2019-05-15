# -*- coding: utf-8 -*-
from __future__ import unicode_literals

import os
import shutil
from datetime import datetime
from doc_to_docx import doc_docx
from doc_2018 import tables_2018
from doc_2017 import tables_2017
from doc_docx_linux import transfer_format
import sqlite3
from docx import Document
import re


def doc_more(document, cursor, date_now, doc_file, doc_more, path_old, doc_year):

    title = ""  # 分发单标题
    next_title = ""
    signer = ""  # 签发
    name = ""  # 分发单名称
    uniquenum = date_now  # 分发单唯一编号
    listnum = ""  # 编号
    auditnum = ""  # 审核编号
    secrecyagreementnum = ""  # 保密责任书编号
    purpose = ""  # 用途
    sendunit = ""  # 发出单位
    sendunitaddr = ""  # 发出单位通讯地址
    sendunitpostcode = ""  # 发出单位邮政编码
    receiveunit = ""  # 接收单位
    receiveunitaddr = ""  # 接收单位通讯地址
    receiveunitpostcode = ""  # 接收单位邮政编码
    handler = ""  # 经办人
    handlerphonenum = ""  # 经办人联系电话(座机)
    handlermobilephonenum = ""  # 经办人联系电话(手机)
    receiver = ""  # 接收人
    receiverphonenum = ""  # 接收人联系电话(座机)
    receivermobilephonenum = ""  # 接收人联系电话(手机)
    sendouttime = ""  # 发出日期
    recievetime = ""  # 接收日期
    selfgetway = False  # 递送方式为自取
    postway = False  # 递送方式为邮寄
    networkway = False  # 递送方式为网络
    sendtoway = False  # 递送方式为送往
    signature = ""  # 递送方式后面的签名
    papermedia = False  # 纸质介质
    cdmedia = False  # 光盘介质
    diskmedia = False  # 硬盘介质
    networkmedia = False  # 网络介质
    othermedia = False  # 其他介质
    medianums = ""  # 介质编号
    mapnums = ""  # 分发单包含的成果的图幅号
    filename = doc_file  # 分发单文件名字
    file = ""  # 分发单文件路径

    remake_str = ""
    is_remake = False
    end_doc = False

    x = 0

    # 读取文本框
    xml = document._element.xml
    pattern = "<v:textbox([\s\S]*?)</v:textbox>"
    regex = re.compile(pattern)
    textbox_list = regex.findall(xml)
    if len(textbox_list) <= 2:
        print("请检查文件{0}内分发单数量".format(doc_file))

    else:
        # 文本框
        data_list = []
        for result in textbox_list:
            pattern2 = "<w:t>(.*?)</w:t>"
            regex2 = re.compile(pattern2)
            textbox_str = regex2.findall(result)
            textbox_str = "".join(textbox_str)
            data_list.append(textbox_str)

        doc_nums = len(textbox_list) / 2
        doc_num = 1
        next_docx = False
        # 读取文本
        for para in document.paragraphs:
            pattern = "[\u4e00-\u9fa5]+"
            regex = re.compile(pattern)
            results = regex.findall(para.text)
            # print(results)
            if results != [] and is_remake is False:
                if "交付清单" in results[0]:
                    title = para.text
                    print("标题：" + results[0])

                if "审批单号" in results:
                    my_num01 = para.text.split(" ")[0].split("：").pop()
                    auditnum = my_num01
                    print("审批单号:" + my_num01)

                if "签发" in results:
                    my_num001 = para.text.split("：").pop().split("签")[0].strip()
                    listnum = my_num001
                    print("编号:" + my_num001)

                if "保密责任证书" in results:
                    my_num02 = para.text.split("：").pop()
                    secrecyagreementnum = my_num02
                    print("保密责任证书:" + my_num02)

                if "递送方式" in results:
                    my_list01 = para.text.split(" ")
                    for my_way in my_list01:
                        if my_way == "自取":
                            selfgetway = True
                        if my_way == "邮寄":
                            postway = True
                        if my_way == "网络":
                            networkway = True
                        if my_way == "送往":
                            sendtoway = True

                if "介质类型" in results:
                    my_list02 = para.text.split(" ")
                    for my_way in my_list02:
                        if my_way == "纸质":
                            selfgetway = True
                        if my_way == "光盘":
                            postway = True
                        if my_way == "硬盘":
                            networkway = True
                        if my_way == "网络":
                            sendtoway = True
                        if my_way == "其他":
                            sendtoway = True

                if "介质编号" in results:
                    my_num03 = para.text.split("：").pop()
                    medianums = my_num03
                    print("介质编号:" + my_num03)

                if "发出单位" in results:
                    my_str01 = para.text.split(" ")[0].split("：").pop()
                    sendunit = my_str01
                    print("发出单位:" + my_str01)

                if "接收单位" in results:
                    my_str02 = para.text.split(" ").pop().split("：").pop()
                    receiveunit = my_str02
                    print("接收单位：" + my_str02)

                if "通讯地址" in results:
                    my_str03 = para.text.split(" ")[0].split("：").pop()
                    sendunitaddr = my_str03
                    print("发出方通讯地址：" + my_str03)
                    my_str03_ = para.text.split("：").pop()
                    receiveunitaddr = my_str03_
                    print("接收方通讯地址：" + receiveunitaddr)

                if "邮政编码" in results:
                    my_str04 = para.text.split(" ")[0].split("：").pop()
                    sendunitpostcode = my_str04
                    print("发出方邮政编码：" + my_str04)

                    my_str04_ = para.text.split("：").pop()
                    receiveunitpostcode = my_str04_
                    print("接收方邮政编码：" + receiveunitpostcode)

                if "座机" in results:
                    my_str05 = para.text.split(" ")[0].split("）").pop()
                    handlerphonenum = my_str05
                    print("发出方联系电话：（座机）：" + my_str05)

                    my_str05_ = para.text.split("）").pop()
                    receiverphonenum = my_str05_
                    print("接收方联系电话：（座机）：" + receiverphonenum)

                if "手机" in results:
                    my_str06 = para.text.split("）")[1].split(" ")[0]
                    handlermobilephonenum = my_str06
                    print("发出方联系电话：（手机）：" + handlermobilephonenum)

                    my_str06_ = para.text.split("）").pop()
                    receivermobilephonenum = my_str06_
                    print("接收方联系电话：（手机）：" + receivermobilephonenum)

                if "发出日期" in results:
                    my_str07 = para.text.split("：")[1].split("接")[0]
                    my_str07 = my_str07.strip()
                    my_str07_year = my_str07.split("年")[0].strip() if my_str07.split("年")[0].strip() != '' else "  "
                    my_str07_mon = my_str07.split("年")[1].split("月")[0].strip() if my_str07.split("年")[1].split("月")[
                                                                                       0].strip() != '' else "  "
                    my_str07_day = my_str07.split("年")[1].split("月")[1].split("日")[0].strip() if \
                    my_str07.split("年")[1].split("月")[1].split("日")[0].strip() != '' else "  "
                    sendouttime = my_str07_year + "年" + my_str07_mon + "月" + my_str07_day + "日"
                    print("发出日期：" + sendouttime)

                    my_str07_ = para.text.split("接收日期：")[1].strip()
                    my_str07_year = my_str07_.split("年")[0].strip() if my_str07_.split("年")[0].strip() != '' else "  "
                    my_str07_mon = my_str07_.split("年")[1].split("月")[0].strip() if my_str07_.split("年")[1].split("月")[
                                                                                        0].strip() != '' else "  "
                    my_str07_day = my_str07_.split("年")[1].split("月")[1].split("日")[0].strip() if \
                    my_str07_.split("年")[1].split("月")[1].split("日")[0].strip() != '' else "  "

                    recievetime = my_str07_year + "年" + my_str07_mon + "月" + my_str07_day + "日"
                    print("接收日期：" + recievetime)

                if "注" in results:
                    if doc_num == doc_nums:

                        end_doc = True
                        my_str09 = para.text.split("：").pop()
                        remake_str = my_str09
                    else:
                        is_remake = True
                        my_str09 = para.text.split("：").pop()
                        remake_str = my_str09
                        continue

            if end_doc == True:
                if "注：" in para.text:
                    pass
                else:
                    remake_str += para.text

            if is_remake:
                if "交付清单" in para.text:
                    is_remake = False
                    next_title = para.text
                    next_docx = True
                else:
                    remake_str += para.text

            if is_remake == False and next_docx ==True:
                name = data_list[doc_num*2 -2]
                purpose = data_list[doc_num*2 - 1]
                mapnums = remake_str
                uniquenum = date_now + "_"  + str(doc_num)
                row = (title, signer, name, uniquenum, listnum, auditnum, secrecyagreementnum, purpose, sendunit, sendunitaddr, sendunitpostcode, receiveunit, receiveunitaddr, receiveunitpostcode, handler, handlerphonenum, handlermobilephonenum, receiver, receiverphonenum, receivermobilephonenum, sendouttime, recievetime, selfgetway, postway, networkway, sendtoway, signature, papermedia, cdmedia, diskmedia, networkmedia, othermedia, medianums, mapnums, filename, file)
                cursor.execute("INSERT INTO results_handoutlist (title, signer, name, uniquenum, listnum, auditnum, secrecyagreementnum, purpose, sendunit, sendunitaddr, sendunitpostcode, receiveunit, receiveunitaddr, receiveunitpostcode, handler, handlerphonenum, handlermobilephonenum, receiver, receiverphonenum, receivermobilephonenum, sendouttime, recievetime, selfgetway, postway, networkway, sendtoway, signature, papermedia, cdmedia, diskmedia, networkmedia, othermedia, medianums, mapnums, filename, file) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", row)
                doc_num += 1
                next_docx = False
                title = next_title

        if end_doc == True:
            name = data_list[doc_num * 2 - 2]
            purpose = data_list[doc_num * 2 - 1]
            mapnums = remake_str
            uniquenum = date_now + "_" + str(doc_num)
            row = (
            title, signer, name, uniquenum, listnum, auditnum, secrecyagreementnum, purpose, sendunit, sendunitaddr,
            sendunitpostcode, receiveunit, receiveunitaddr, receiveunitpostcode, handler, handlerphonenum,
            handlermobilephonenum, receiver, receiverphonenum, receivermobilephonenum, sendouttime, recievetime,
            selfgetway, postway, networkway, sendtoway, signature, papermedia, cdmedia, diskmedia, networkmedia,
            othermedia, medianums, mapnums, filename, file)
            cursor.execute(
                "INSERT INTO results_handoutlist (title, signer, name, uniquenum, listnum, auditnum, secrecyagreementnum, purpose, sendunit, sendunitaddr, sendunitpostcode, receiveunit, receiveunitaddr, receiveunitpostcode, handler, handlerphonenum, handlermobilephonenum, receiver, receiverphonenum, receivermobilephonenum, sendouttime, recievetime, selfgetway, postway, networkway, sendtoway, signature, papermedia, cdmedia, diskmedia, networkmedia, othermedia, medianums, mapnums, filename, file) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                row)


        # 读取表格
        name = ""  # 成果资料名称
        secretlevel = ""  # 成果秘密级别
        resultnum = ""  # 成果数量
        datasize = ""  # 成果数据量
        formatormedia = ""  # 格式/介质
        remarks = ""  # 成果资料备注
        handoutlist_uniquenum = ""  # 分发单唯一编号
        if doc_year == 2017:
            p = 1
            table_nums = 1
            tables = document.tables
            for table in tables:
                for i in table.rows:
                    table_list = []
                    for x in i.cells:
                        table_list.append(x.text)
                    print(table_list)
                    name = table_list[1]
                    secretlevel = table_list[2]
                    resultnum = table_list[3]
                    datasize = ""
                    formatormedia = table_list[4]
                    remarks = table_list[5]
                    handoutlist_uniquenum = date_now

                    if p == 1:
                        p = 2
                    else:
                        row2 = (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum)
                        cursor.execute(
                            "INSERT INTO results_fileinfo (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum) VALUES (?, ?, ?, ?, ?, ?, ?)",
                            row2)
                p = 1
                table_nums += 1

        else:
            p = 1
            table_nums = 1
            tables = document.tables
            for table in tables:
                for i in table.rows:
                    table_list = []
                    for x in i.cells:
                        table_list.append(x.text)
                    name = table_list[1]
                    secretlevel = table_list[2]
                    resultnum = table_list[3]
                    datasize = table_list[4]
                    formatormedia = table_list[5]
                    remarks = table_list[6]
                    handoutlist_uniquenum = date_now + "_" + str(table_nums)

                    if p == 1:
                        p = 2
                    else:
                        print(table_list)
                        row2 = (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum)
                        cursor.execute(
                            "INSERT INTO results_fileinfo (name, secretlevel, resultnum, datasize, formatormedia, remarks, handoutlist_uniquenum) VALUES (?, ?, ?, ?, ?, ?, ?)",
                            row2)
                p = 1
                table_nums += 1

        # shutil.move(path_old, doc_more)

    return date_now, doc_file




if __name__ == '__main__':
    pass