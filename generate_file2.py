# -*- coding: utf-8 -*-
import os
if not os.environ.get("DJANGO_SETTINGS_MODULE"):
    os.environ.setdefault("DJANGO_SETTINGS_MODULE", "ResultManage.settings")
import django
django.setup()
import time
import random
import subprocess
import docx
from datetime import datetime
import re
# import psycopg2
import sqlite3
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION


def generate_docx(handoutlist, fileinfo_list, templates_dir, handoutlist_docxs, handoutlist_docs):
    result_count = len(fileinfo_list)
    template_filepath = os.path.join(templates_dir, "template" + str(result_count) + ".docx")
    doc = docx.Document(template_filepath)
    # 如果只为null,替换成空字符串""
    handoutlist_propertys = [
        handoutlist.title,
        handoutlist.signer,
        handoutlist.name,
        handoutlist.listnum,
        handoutlist.auditnum,
        handoutlist.secrecyagreementnum,
        handoutlist.purpose,
        handoutlist.sendunit,
        handoutlist.sendunitaddr,
        handoutlist.sendunitpostcode,
        handoutlist.receiveunit,
        handoutlist.receiveunitaddr,
        handoutlist.receiveunitpostcode,
        handoutlist.handler,
        handoutlist.handlerphonenum,
        handoutlist.handlermobilephonenum,
        handoutlist.receiver,
        handoutlist.receiverphonenum,
        handoutlist.receivermobilephonenum,
        handoutlist.sendouttime,
        handoutlist.recievetime,
        handoutlist.signature,
        handoutlist.medianums,
        handoutlist.mapnums,
        handoutlist.filename,
        handoutlist.file,
    ]
    # 防止有的值为null出错
    for handoutlist_property in handoutlist_propertys:
        if handoutlist_property is None or handoutlist_property == "None":
            handoutlist_propertys[handoutlist_propertys.index(handoutlist_property)] = ""

    # print(len(doc.paragraphs))
    # 清单名称
    title = handoutlist_propertys[0]
    # doc.paragraphs[0].runs[0].text = title
    title_runs = doc.paragraphs[0].runs
    for run in title_runs:
        if title_runs.index(run) == 0:
            run.text = title
        else:
            run.text = ""
    # 编号
    listnum = handoutlist_propertys[3]
    signer = handoutlist_propertys[0]
    # doc.paragraphs[1].text = doc.paragraphs[1].text.replace("listnum", listnum)
    # 签发
    # doc.paragraphs[1].text = doc.paragraphs[1].text.replace("signer", signer)
    # print(doc.paragraphs[1].text)
    # print(doc.paragraphs[1].text)
    # for run in doc.paragraphs[1].runs:
    #     print(run.text)
    #     if run.text in listnum or run.text in signer:
    #         run.underline = True
    # print(doc.paragraphs[1].runs[3].text)

    for run in doc.paragraphs[1].runs:
        if run.text.strip() == "listnum":
            run.text = listnum
            run.underline = True
        if run.text.strip() == "signer":
            run.text = signer
            run.underline = True

    # 名称
    name = handoutlist.name
    doc.paragraphs[3].text = doc.paragraphs[3].text.replace("name", name)

    # 用途
    perpose = handoutlist.purpose
    doc.paragraphs[4].text = doc.paragraphs[4].text.replace("purpose", perpose)
    # 审单
    auditnum = handoutlist.auditnum
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("auditnum", auditnum)
    # 协议编号
    secrecyagreementnum = handoutlist.secrecyagreementnum
    doc.paragraphs[5].text = doc.paragraphs[5].text.replace("secrecyagreementnum", secrecyagreementnum)
    # for run in doc.paragraphs[5].runs:
    #     if run.text == "支援保障":
    #         run.text = perpose
        # if run.text == "secrecyagreementnum":
        #     run.text = secrecyagreementnum

    # 递送方式, 递送签名
    # selfgetway  postway  networkway  sendtoway

    selfgetway = handoutlist.selfgetway
    postway = handoutlist.postway
    networkway = handoutlist.networkway
    sendtoway = handoutlist.sendtoway
    signature = handoutlist.signature
    way1 = "☑自取" if selfgetway else "□自取"
    way2 = "☑邮寄" if postway else "□邮寄"
    way3 = "☑网络" if networkway else "□网络"
    way4 = "☑送往" if sendtoway else "□送往"
    ways = way1 + "  " + way2 + "  " + way3 + "  " + way4
    # doc.paragraphs[6].text = doc.paragraphs[6].text.replace("□自取  □邮寄  □网络  □送往 （签名）signature",ways)
    doc.paragraphs[6].text = "递送方式： " + ways + " （签名）"
    run6 = doc.paragraphs[6].add_run(signature)
    run6.underline = True

    # 介质类型
    papermedia = handoutlist.papermedia
    cdmedia = handoutlist.cdmedia
    diskmedia = handoutlist.diskmedia
    networkmedia = handoutlist.networkmedia
    othermedia = handoutlist.othermedia
    media1 = "☑纸质" if papermedia else "□纸质"
    media2 = "☑光盘" if cdmedia else "□光盘"
    media3 = "☑硬盘" if diskmedia else "□硬盘"
    media4 = "☑网络" if networkmedia else "□网络"
    media5 = "☑其他" if othermedia else "□其他"
    # 介质类型
    medias = media1 + "  " + media2 + "  " + media3 + "  " + media4 + "  " + media5
    # 介质编号
    medianums = handoutlist.medianums
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("□纸质  □光盘  □硬盘  □网络  □其他", medias)
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace("medianums", medianums)
    doc.paragraphs[7].text = "介质类型： " + medias + "      介质编号：" + medianums

    # 发出单位
    sendunit = handoutlist.sendunit
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("sendunit", sendunit)
    # 接收单位
    receiveunit = handoutlist.receiveunit
    doc.paragraphs[9].text = doc.paragraphs[9].text.replace("receiveunit", receiveunit)

    # 发出单位通讯地址
    sendunitaddr = handoutlist.sendunitaddr
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("sendunitaddr", sendunitaddr)
    # 接收单位通讯地址
    receiveunitaddr = handoutlist.receiveunitaddr
    doc.paragraphs[12].text = doc.paragraphs[12].text.replace("receiveunitaddr", receiveunitaddr)
    # 发出单位邮政编码
    sendunitpostcode = handoutlist.sendunitpostcode
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("sendunitpostcode", sendunitpostcode)
    # 接收单位邮政编码
    receiveunitpostcode = handoutlist.receiveunitpostcode
    doc.paragraphs[13].text = doc.paragraphs[13].text.replace("receiveunitpostcode", receiveunitpostcode)

    # 经办人
    handler = handoutlist.handler
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("handler", handler)
    # 接收人
    receiver = handoutlist.receiver
    doc.paragraphs[14].text = doc.paragraphs[14].text.replace("receiver", receiver)
    # 经办人联系电话(座机)
    handlerphonenum = handoutlist.handlerphonenum
    if handlerphonenum is None:
        handlerphonenum =""
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("handlerphonenum", handlerphonenum)
    # 接收人联系电话(座机)
    receiverphonenum = handoutlist.receiverphonenum
    doc.paragraphs[15].text = doc.paragraphs[15].text.replace("receiverphonenum", receiverphonenum)
    # 经办人联系电话(手机)
    handlermobilephonenum = handoutlist.handlermobilephonenum
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("handlermobilephonenum", handlermobilephonenum)
    # 接收人联系电话(手机)
    receivermobilephonenum = handoutlist.receivermobilephonenum
    doc.paragraphs[16].text = doc.paragraphs[16].text.replace("receivermobilephonenum", receivermobilephonenum)

    # 发出日期
    sendouttime = handoutlist.sendouttime
    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("sendouttime", sendouttime)
    # 接收日期
    recievetime = handoutlist.recievetime
    doc.paragraphs[17].text = doc.paragraphs[17].text.replace("recievetime", recievetime)

    table = doc.tables[0]
    count = 1
    for fileinfo in fileinfo_list:
        row = table.rows[count]
        # row_cells = row.cells
        # for cell in row_cells:
        #     # 表格内容居中
        #     # cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
        #     if row_cells.index(cell) == 1:
        #         cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平左对齐
        #     # else:
        #     #     cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 水平居中对齐
        #     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # 垂直对齐
        # print(len(row.cells[1].paragraphs[0].runs))

        # row.cells[0].paragraphs[0].runs[0].text = str(count)
        # row.cells[1].paragraphs[0].runs[0].text = fileinfo["name"]
        # row.cells[2].paragraphs[0].runs[0].text = fileinfo["secretlevel"]
        # row.cells[3].paragraphs[0].runs[0].text = fileinfo["resultnum"]
        # row.cells[4].paragraphs[0].runs[0].text = fileinfo["datasize"]
        # row.cells[5].paragraphs[0].runs[0].text = fileinfo["formatormedia"]
        # row.cells[6].paragraphs[0].runs[0].text = fileinfo["remarks"]

        row.cells[0].paragraphs[0].runs[0].text = str(count)
        # name
        row.cells[1].paragraphs[0].runs[0].text = fileinfo.name
        # secretlevel
        row.cells[2].paragraphs[0].runs[0].text = fileinfo.secretlevel
        # resultnum
        row.cells[3].paragraphs[0].runs[0].text = fileinfo.resultnum
        # datasize
        row.cells[4].paragraphs[0].runs[0].text = fileinfo.datasize
        # formatormedium
        row.cells[5].paragraphs[0].runs[0].text = fileinfo.formatormedium
        # remarks
        row.cells[6].paragraphs[0].runs[0].text = fileinfo.remarks
        count += 1

    mapnums = handoutlist.mapnums
    # doc.paragraphs[18].text = doc.paragraphs[18].text.replace("mapnums", mapnums)
    doc.paragraphs[18].text = doc.paragraphs[18].text + mapnums
    # 添加段落
    # doc.add_paragraph(mapnums)
    # filename = handoutlist[34]
    filename = "测绘" + datetime.now().strftime("%Y") + "-" + "%04d" % handoutlist.id + ".docx"
    docx_filepath = os.path.join(handoutlist_docxs, filename)
    doc.save(docx_filepath)

    handoutlist.file = os.path.join("handoutlist_docxs",filename)
    handoutlist.save()
    print("调用了generate_docx函数")
    return handoutlist

    # 转换为doc格式
    # subprocess.check_output(
    #     ["soffice", "--headless", "--invisible", "--convert-to", "doc", docx_filepath, "--outdir", handoutlist_docs])


if __name__ == '__main__':
    # generate_docx(handoutlist, fileinfo_list, templates_dir, handoutlist_docxs, handoutlist_docs)
    pass